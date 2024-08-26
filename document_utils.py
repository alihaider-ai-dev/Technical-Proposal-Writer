from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def process_document(input_path, output_path, apply_rtl=False):
    # Load the document
    doc = Document(input_path)

    def add_borders_to_table(table):
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblBorders = tblPr.find(qn('w:tblBorders'))
        if tblBorders is None:
            tblBorders = OxmlElement('w:tblBorders')
            tblPr.append(tblBorders)
        for border_name in [
                'top', 'left', 'bottom', 'right', 'insideH', 'insideV'
        ]:
            border = tblBorders.find(qn(f'w:{border_name}'))
            if border is None:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:space'), '0')
                border.set(qn('w:sz'),
                           '4')  # Border size (e.g., '4' for 1pt border)
                border.set(qn('w:space'), '0')
                border.set(qn('w:space'), '0')
                tblBorders.append(border)
            else:
                border.set(qn('w:space'), '0')
                border.set(qn('w:sz'), '4')  # Ensure border size

    def set_rtl_paragraph(paragraph):
        p = paragraph._element
        pPr = p.get_or_add_pPr()
        pBidi = pPr.find(qn('w:bidi'))
        if pBidi is None:
            pBidi = OxmlElement('w:bidi')
            pPr.append(pBidi)
        pJc = pPr.find(qn('w:jc'))
        if pJc is None:
            pJc = OxmlElement('w:jc')
            pPr.append(pJc)
        pJc.set(qn('w:val'), 'right')

    def set_rtl_table(table):
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblBidi = tblPr.find(qn('w:tblBidi'))
        if tblBidi is None:
            tblBidi = OxmlElement('w:tblBidi')
            tblPr.append(tblBidi)

    def remove_extra_page_breaks():
        # Remove extra page breaks by checking paragraphs
        for para in doc.paragraphs:
            if para.text.strip() == "":
                # Remove empty paragraphs that might be extra page breaks
                p = para._element
                p.getparent().remove(p)

    # Apply borders to all tables
    for table in doc.tables:
        add_borders_to_table(table)

    # Apply RTL if needed
    if apply_rtl:
        for para in doc.paragraphs:
            set_rtl_paragraph(para)
        for table in doc.tables:
            set_rtl_table(table)

    # Remove extra page breaks
    remove_extra_page_breaks()

    # Save the modified document
    doc.save(output_path)
